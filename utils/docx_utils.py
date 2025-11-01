import re
from collections import OrderedDict
from docx import Document

PLACEHOLDER_RE = re.compile(r"\[[^\[\]]+\]")  # e.g., [Company Name]

def _iter_paragraphs(element):
    """Yield all paragraphs within a document element, including nested tables."""
    if hasattr(element, "paragraphs"):
        for paragraph in element.paragraphs:
            yield paragraph
    if hasattr(element, "tables"):
        for table in element.tables:
            for row in table.rows:
                for cell in row.cells:
                    yield from _iter_paragraphs(cell)

def _apply_mapping(text: str, mapping: dict[str, str | None]) -> str:
    updated = text
    for placeholder, value in mapping.items():
        if not placeholder or value is None:
            continue
        replacement = str(value)
        updated = updated.replace(placeholder, replacement)
    return updated

def extract_placeholders(doc_path: str) -> list[str]:
    doc = Document(doc_path)
    ordered = OrderedDict()
    for paragraph in _iter_paragraphs(doc):
        text = paragraph.text or ""
        for match in PLACEHOLDER_RE.findall(text):
            ordered.setdefault(match.strip(), None)
    return list(ordered.keys())

def replace_placeholders(doc_path: str, mapping: dict[str, str | None], out_path: str) -> None:
    doc = Document(doc_path)
    if not mapping:
        doc.save(out_path)
        return

    for paragraph in _iter_paragraphs(doc):
        original = paragraph.text or ""
        if not original:
            continue
        updated = _apply_mapping(original, mapping)
        if updated != original:
            paragraph.text = updated

    doc.save(out_path)

def build_preview_text(doc_path: str, mapping: dict[str, str | None]) -> str:
    doc = Document(doc_path)
    buf = []
    for paragraph in _iter_paragraphs(doc):
        text = paragraph.text or ""
        updated = _apply_mapping(text, mapping)
        if updated.strip():
            buf.append(updated)
    return "\n\n".join(buf)
